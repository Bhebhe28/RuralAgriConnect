from pathlib import Path
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
source = ROOT / "Security_Assignment_NIST_CSF2_OWASP2025_RuralAgriConnect.md"
target = ROOT / "Security_Assignment_NIST_CSF2_OWASP2025_RuralAgriConnect.docx"

doc = Document()

for raw_line in source.read_text(encoding="utf-8").splitlines():
    line = raw_line.rstrip()
    if not line:
        doc.add_paragraph("")
        continue

    if line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=3)
    elif line.startswith("- "):
        doc.add_paragraph(line[2:].strip(), style="List Bullet")
    else:
        # Keep markdown tables/readable blocks as plain paragraphs for Word compatibility
        doc.add_paragraph(line)

doc.save(target)
print(f"Generated: {target}")
